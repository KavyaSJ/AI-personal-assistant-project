
import re
from datetime import datetime, timedelta

import docx
import pdfplumber


DEFAULT_CLASS_DURATION_MINUTES = 60
PDF_LOW_TEXT_THRESHOLD = 80
OCR_PAGE_LIMIT = 20


# -----------------------------
# Parsing layer
# -----------------------------

# Cleans repeated whitespace into a single readable space.
def _clean_space(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()


# Splits a text blob into readable block units.
def _split_into_blocks(text: str) -> list:
    if not text:
        return []

    blocks = []
    for chunk in re.split(r"\n\s*\n", text):
        cleaned = chunk.strip()
        if cleaned:
            blocks.append(cleaned)

    if blocks:
        return blocks

    return [line.strip() for line in text.splitlines() if line.strip()]


# Converts extracted table data into a compact list of row dicts.
def _normalize_table_rows(table) -> list:
    rows = []
    for row in table or []:
        rows.append([_clean_space(cell or "") for cell in (row or [])])
    return rows


# Attempts OCR on PDFs only when text extraction is weak.
def _run_optional_pdf_ocr(file_path: str, max_pages: int = OCR_PAGE_LIMIT) -> dict:
    result = {
        "available": False,
        "attempted": False,
        "used": False,
        "status": "not_attempted",
        "pages": [],
        "text": "",
        "message": "",
    }

    try:
        from pdf2image import convert_from_path
        import pytesseract
    except Exception as e:
        result["status"] = "dependencies_missing"
        result["message"] = str(e)
        return result

    result["available"] = True
    result["attempted"] = True

    try:
        images = convert_from_path(file_path, first_page=1, last_page=max_pages)
        ocr_pages = []

        for index, image in enumerate(images, start=1):
            page_text = pytesseract.image_to_string(image) or ""
            cleaned = page_text.strip()
            ocr_pages.append({
                "page_number": index,
                "text": cleaned,
            })

        result["pages"] = ocr_pages
        result["text"] = "\n\n".join([item["text"] for item in ocr_pages if item["text"]]).strip()
        result["used"] = bool(result["text"])
        result["status"] = "success" if result["used"] else "empty"
        return result
    except Exception as e:
        result["status"] = "failed"
        result["message"] = str(e)
        return result


# Parses a PDF into structured content with pages, blocks, tables, and quality metadata.
def parse_pdf_document(file_path: str, enable_ocr: bool = True) -> dict:
    parsed = {
        "source_type": "pdf",
        "raw_text": "",
        "pages": [],
        "blocks": [],
        "tables": [],
        "quality": {
            "parser": "pdfplumber",
            "page_count": 0,
            "nonempty_pages": 0,
            "char_count": 0,
            "word_count": 0,
            "table_count": 0,
            "quality_score": 0.0,
            "quality_flags": [],
            "ocr_attempted": False,
            "ocr_used": False,
            "ocr_status": "not_attempted",
            "ocr_message": "",
            "weak_text_extraction": False,
        },
    }

    if not file_path:
        parsed["quality"]["quality_flags"].append("missing_file_path")
        return parsed

    page_texts = []

    try:
        with pdfplumber.open(file_path) as pdf:
            parsed["quality"]["page_count"] = len(pdf.pages)

            for page_index, page in enumerate(pdf.pages, start=1):
                page_text = (page.extract_text() or "").strip()
                page_tables = []

                for table_index, table in enumerate(page.extract_tables() or [], start=1):
                    clean_table = _normalize_table_rows(table)
                    row_count = len(clean_table)
                    col_count = max((len(row) for row in clean_table), default=0)

                    if row_count >= 2 and col_count >= 2:
                        table_obj = {
                            "page_number": page_index,
                            "table_index": table_index,
                            "rows": clean_table,
                        }
                        page_tables.append(table_obj)
                        parsed["tables"].append(table_obj)

                blocks = []
                for block_index, block_text in enumerate(_split_into_blocks(page_text), start=1):
                    block_obj = {
                        "page_number": page_index,
                        "block_index": block_index,
                        "text": block_text,
                    }
                    blocks.append(block_obj)
                    parsed["blocks"].append(block_obj)

                parsed["pages"].append({
                    "page_number": page_index,
                    "text": page_text,
                    "blocks": blocks,
                    "tables": page_tables,
                })

                if page_text:
                    parsed["quality"]["nonempty_pages"] += 1
                    page_texts.append(page_text)
    except Exception as e:
        parsed["quality"]["quality_flags"].append("pdf_parse_failed")
        parsed["quality"]["ocr_message"] = str(e)
        return parsed

    parsed["raw_text"] = "\n\n".join(page_texts).strip()

    if len(parsed["raw_text"]) < PDF_LOW_TEXT_THRESHOLD:
        parsed["quality"]["weak_text_extraction"] = True
        parsed["quality"]["quality_flags"].append("low_text_volume")

        if enable_ocr:
            ocr_result = _run_optional_pdf_ocr(file_path)
            parsed["quality"]["ocr_attempted"] = ocr_result.get("attempted", False)
            parsed["quality"]["ocr_used"] = ocr_result.get("used", False)
            parsed["quality"]["ocr_status"] = ocr_result.get("status", "not_attempted")
            parsed["quality"]["ocr_message"] = ocr_result.get("message", "")

            if ocr_result.get("used"):
                ocr_pages = {item["page_number"]: item.get("text", "") for item in ocr_result.get("pages", [])}

                page_texts_after_ocr = []
                parsed["blocks"] = []
                for page in parsed["pages"]:
                    page_number = page["page_number"]
                    if not page.get("text") and ocr_pages.get(page_number):
                        page["text"] = ocr_pages[page_number].strip()

                    page["blocks"] = []
                    for block_index, block_text in enumerate(_split_into_blocks(page.get("text", "")), start=1):
                        block_obj = {
                            "page_number": page_number,
                            "block_index": block_index,
                            "text": block_text,
                        }
                        page["blocks"].append(block_obj)
                        parsed["blocks"].append(block_obj)

                    if page.get("text"):
                        page_texts_after_ocr.append(page["text"])

                parsed["raw_text"] = "\n\n".join(page_texts_after_ocr).strip()
                if parsed["raw_text"]:
                    parsed["quality"]["nonempty_pages"] = sum(1 for page in parsed["pages"] if page.get("text"))
                    if "low_text_volume" in parsed["quality"]["quality_flags"] and len(parsed["raw_text"]) >= PDF_LOW_TEXT_THRESHOLD:
                        parsed["quality"]["quality_flags"].remove("low_text_volume")
                        parsed["quality"]["weak_text_extraction"] = False

    parsed["quality"]["table_count"] = len(parsed["tables"])
    parsed["quality"]["char_count"] = len(parsed["raw_text"])
    parsed["quality"]["word_count"] = len(parsed["raw_text"].split())

    page_score = 0.0
    if parsed["quality"]["page_count"]:
        page_score = parsed["quality"]["nonempty_pages"] / parsed["quality"]["page_count"]

    text_score = 1.0 if parsed["quality"]["char_count"] >= 1000 else min(parsed["quality"]["char_count"] / 1000.0, 1.0)
    table_bonus = 0.1 if parsed["quality"]["table_count"] else 0.0

    quality_score = round(min(page_score * 0.6 + text_score * 0.3 + table_bonus, 1.0), 2)
    parsed["quality"]["quality_score"] = quality_score

    if not parsed["raw_text"]:
        parsed["quality"]["quality_flags"].append("no_text_extracted")
    if not parsed["tables"]:
        parsed["quality"]["quality_flags"].append("no_tables_found")

    return parsed


# Parses a DOCX into structured content with pseudo-pages, blocks, tables, and quality metadata.
def parse_docx_document(file_path: str) -> dict:
    parsed = {
        "source_type": "docx",
        "raw_text": "",
        "pages": [],
        "blocks": [],
        "tables": [],
        "quality": {
            "parser": "python-docx",
            "page_count": 1,
            "nonempty_pages": 0,
            "char_count": 0,
            "word_count": 0,
            "table_count": 0,
            "quality_score": 0.0,
            "quality_flags": [],
            "ocr_attempted": False,
            "ocr_used": False,
            "ocr_status": "not_applicable",
            "ocr_message": "",
            "weak_text_extraction": False,
        },
    }

    if not file_path:
        parsed["quality"]["quality_flags"].append("missing_file_path")
        return parsed

    blocks = []
    tables = []
    text_parts = []

    try:
        document = docx.Document(file_path)

        for paragraph in document.paragraphs:
            paragraph_text = (paragraph.text or "").strip()
            if paragraph_text:
                text_parts.append(paragraph_text)
                blocks.append({
                    "page_number": 1,
                    "block_index": len(blocks) + 1,
                    "text": paragraph_text,
                    "block_type": "paragraph",
                })

        for table_index, table in enumerate(document.tables, start=1):
            clean_table = []
            for row in table.rows:
                clean_row = []
                for cell in row.cells:
                    cell_text = _clean_space(cell.text or "")
                    clean_row.append(cell_text)
                clean_table.append(clean_row)

            row_count = len(clean_table)
            col_count = max((len(row) for row in clean_table), default=0)
            if row_count >= 2 and col_count >= 2:
                table_obj = {
                    "page_number": 1,
                    "table_index": table_index,
                    "rows": clean_table,
                }
                tables.append(table_obj)
                blocks.append({
                    "page_number": 1,
                    "block_index": len(blocks) + 1,
                    "text": "\n".join([" | ".join(row) for row in clean_table if any(row)]).strip(),
                    "block_type": "table",
                })
                for row in clean_table:
                    row_text = " | ".join([cell for cell in row if cell]).strip()
                    if row_text:
                        text_parts.append(row_text)
    except Exception as e:
        parsed["quality"]["quality_flags"].append("docx_parse_failed")
        parsed["quality"]["ocr_message"] = str(e)
        return parsed

    parsed["raw_text"] = "\n\n".join([part for part in text_parts if part]).strip()
    parsed["tables"] = tables
    parsed["blocks"] = blocks
    parsed["pages"] = [{
        "page_number": 1,
        "text": parsed["raw_text"],
        "blocks": blocks,
        "tables": tables,
    }]

    parsed["quality"]["nonempty_pages"] = 1 if parsed["raw_text"] else 0
    parsed["quality"]["table_count"] = len(parsed["tables"])
    parsed["quality"]["char_count"] = len(parsed["raw_text"])
    parsed["quality"]["word_count"] = len(parsed["raw_text"].split())
    parsed["quality"]["quality_score"] = round(
        min((1.0 if parsed["raw_text"] else 0.0) * 0.85 + (0.15 if parsed["tables"] else 0.0), 1.0),
        2,
    )

    if len(parsed["raw_text"]) < PDF_LOW_TEXT_THRESHOLD:
        parsed["quality"]["weak_text_extraction"] = True
        parsed["quality"]["quality_flags"].append("low_text_volume")
    if not parsed["tables"]:
        parsed["quality"]["quality_flags"].append("no_tables_found")

    return parsed


# Extracts text from a PDF file.
def extract_text_from_pdf(file_path):
    return parse_pdf_document(file_path).get("raw_text", "")


# Extracts text from a DOCX file, including table cells.
def extract_text_from_docx(file_path):
    return parse_docx_document(file_path).get("raw_text", "")


# Extracts native tables from a PDF file.
def extract_tables_from_pdf(file_path: str) -> list:
    return parse_pdf_document(file_path).get("tables", [])


# Extracts native tables from a DOCX file.
def extract_tables_from_docx(file_path: str) -> list:
    return parse_docx_document(file_path).get("tables", [])


# -----------------------------
# Detection and normalization helpers
# -----------------------------

# Detects whether a document looks like a course outline, invitation, or general document.
def detect_document_type(text: str) -> str:
    if not text:
        return "general"

    text_lower = text.lower()

    course_keywords = [
        "course outline",
        "syllabus",
        "week 1",
        "lecture 1",
        "session 1",
        "midterm",
        "final exam",
        "assignment",
        "course code",
        "course title",
        "office hours",
        "tutorial",
        "quiz",
        "lab section",
        "textbook",
    ]

    invitation_keywords = [
        "you are invited",
        "cordially invited",
        "please join us",
        "join us for",
        "invite you to",
        "rsvp",
        "kindly rsvp",
        "the pleasure of your company",
        "wedding",
        "reception",
        "birthday party",
        "baby shower",
        "graduation ceremony",
        "housewarming",
        "annual gala",
        "celebration of",
    ]

    course_score = sum(1 for keyword in course_keywords if keyword in text_lower)
    invitation_score = sum(1 for keyword in invitation_keywords if keyword in text_lower)

    if course_score >= 2 or ("assignment" in text_lower and ("lecture" in text_lower or "midterm" in text_lower)):
        return "course_outline"

    if invitation_score >= 2 or ("rsvp" in text_lower and ("venue" in text_lower or "join us" in text_lower)):
        return "invitation"

    return "general"


# Returns a small context window around the first keyword match.
def _get_context_window(text: str, keyword: str, window: int = 200) -> str:
    if not text or not keyword:
        return ""

    match = re.search(re.escape(keyword), text, flags=re.IGNORECASE)
    if not match:
        return ""

    start = max(0, match.start() - window)
    end = min(len(text), match.end() + window)
    return text[start:end].strip()


# Safely converts a date string into YYYY-MM-DD.
def _parse_date_string(date_text: str, default_year=None) -> str:
    if not date_text:
        return ""

    cleaned = _clean_space(date_text)
    cleaned = cleaned.replace("Sept ", "Sep ")
    cleaned = re.sub(r"(\d)(st|nd|rd|th)\b", r"\1", cleaned, flags=re.IGNORECASE)
    cleaned = cleaned.replace("–", "-").replace("—", "-")
    cleaned = re.sub(r"\b(on|by|due|from|to|until|through)\b", " ", cleaned, flags=re.IGNORECASE)
    cleaned = _clean_space(cleaned).strip(",.-")

    year_to_use = default_year or datetime.now().year

    date_formats = [
        "%Y-%m-%d",
        "%B %d, %Y",
        "%b %d, %Y",
        "%d %B %Y",
        "%d %b %Y",
        "%B %d %Y",
        "%b %d %Y",
        "%d/%m/%Y",
        "%m/%d/%Y",
        "%d-%m-%Y",
        "%m-%d-%Y",
        "%B %d",
        "%b %d",
        "%d %B",
        "%d %b",
        "%m/%d",
        "%d/%m",
    ]

    for fmt in date_formats:
        try:
            parsed = datetime.strptime(cleaned, fmt)
            if "%Y" not in fmt:
                parsed = parsed.replace(year=year_to_use)
            return parsed.strftime("%Y-%m-%d")
        except ValueError:
            continue

    match = re.search(
        r"\b([A-Za-z]{3,9}\.?\s+\d{1,2})(?:,?\s+(\d{4}))?\b",
        cleaned,
        flags=re.IGNORECASE,
    )
    if match:
        month_day = match.group(1).replace(".", "")
        year = int(match.group(2)) if match.group(2) else year_to_use
        for fmt in ["%B %d %Y", "%b %d %Y"]:
            try:
                parsed = datetime.strptime(f"{month_day} {year}", fmt)
                return parsed.strftime("%Y-%m-%d")
            except ValueError:
                continue

    slash_match = re.search(r"\b(\d{1,2})[/-](\d{1,2})(?:[/-](\d{2,4}))?\b", cleaned)
    if slash_match:
        first = int(slash_match.group(1))
        second = int(slash_match.group(2))
        year_part = slash_match.group(3)
        if year_part:
            year = int(year_part)
            if year < 100:
                year += 2000
        else:
            year = year_to_use

        candidates = [
            (first, second),
            (second, first),
        ]
        for month, day in candidates:
            try:
                parsed = datetime(year, month, day)
                return parsed.strftime("%Y-%m-%d")
            except ValueError:
                continue

    return ""


# Safely converts a time string into HH:MM 24-hour format.
def _parse_time_string(time_text: str) -> str:
    if not time_text:
        return ""

    cleaned = _clean_space(time_text)
    cleaned = cleaned.replace("a.m.", "AM").replace("p.m.", "PM")
    cleaned = cleaned.replace("a.m", "AM").replace("p.m", "PM")
    cleaned = cleaned.replace("am", "AM").replace("pm", "PM")
    cleaned = cleaned.replace(" noon", " 12:00 PM")
    cleaned = cleaned.replace(" midnight", " 12:00 AM")
    cleaned = _clean_space(cleaned)

    formats = [
        "%H:%M",
        "%H%M",
        "%I:%M %p",
        "%I %p",
        "%I:%M%p",
        "%I%p",
    ]

    for fmt in formats:
        try:
            parsed = datetime.strptime(cleaned, fmt)
            return parsed.strftime("%H:%M")
        except ValueError:
            continue

    match = re.search(r"\b(\d{1,2})(?::(\d{2}))?\s*(AM|PM)\b", cleaned, flags=re.IGNORECASE)
    if match:
        hour = int(match.group(1))
        minute = int(match.group(2) or 0)
        meridiem = match.group(3).upper()
        try:
            parsed = datetime.strptime(f"{hour}:{minute:02d} {meridiem}", "%I:%M %p")
            return parsed.strftime("%H:%M")
        except ValueError:
            return ""

    match24 = re.search(r"\b([01]?\d|2[0-3]):([0-5]\d)\b", cleaned)
    if match24:
        return f"{int(match24.group(1)):02d}:{int(match24.group(2)):02d}"

    return ""


# Returns an end time after adding a default class duration.
def _infer_end_time(start_time: str, minutes: int = DEFAULT_CLASS_DURATION_MINUTES) -> str:
    if not start_time:
        return ""

    try:
        start_dt = datetime.strptime(start_time, "%H:%M")
        end_dt = start_dt + timedelta(minutes=minutes)
        return end_dt.strftime("%H:%M")
    except ValueError:
        return ""


# Normalizes weekday abbreviations into full weekday names.
def _normalize_weekday_tokens(tokens) -> list:
    if not tokens:
        return []

    if isinstance(tokens, str):
        tokens = [tokens]

    mapping = {
        "m": ["Monday"],
        "mo": ["Monday"],
        "mon": ["Monday"],
        "monday": ["Monday"],
        "t": ["Tuesday"],
        "tu": ["Tuesday"],
        "tue": ["Tuesday"],
        "tues": ["Tuesday"],
        "tuesday": ["Tuesday"],
        "w": ["Wednesday"],
        "we": ["Wednesday"],
        "wed": ["Wednesday"],
        "wednesday": ["Wednesday"],
        "th": ["Thursday"],
        "thu": ["Thursday"],
        "thur": ["Thursday"],
        "thurs": ["Thursday"],
        "thursday": ["Thursday"],
        "r": ["Thursday"],
        "f": ["Friday"],
        "fr": ["Friday"],
        "fri": ["Friday"],
        "friday": ["Friday"],
        "sa": ["Saturday"],
        "sat": ["Saturday"],
        "saturday": ["Saturday"],
        "su": ["Sunday"],
        "sun": ["Sunday"],
        "sunday": ["Sunday"],
        "tr": ["Tuesday", "Thursday"],
        "mw": ["Monday", "Wednesday"],
        "wf": ["Wednesday", "Friday"],
        "mwf": ["Monday", "Wednesday", "Friday"],
        "tth": ["Tuesday", "Thursday"],
        "tu th": ["Tuesday", "Thursday"],
        "tue thu": ["Tuesday", "Thursday"],
    }

    ordered_days = [
        "Monday",
        "Tuesday",
        "Wednesday",
        "Thursday",
        "Friday",
        "Saturday",
        "Sunday",
    ]

    found = []

    for token in tokens:
        if token is None:
            continue

        text = _clean_space(str(token)).lower().strip(",;|")
        if not text:
            continue

        if text in mapping:
            found.extend(mapping[text])
            continue

        split_parts = re.split(r"[,/&]|\band\b", text)
        if len(split_parts) > 1:
            for part in split_parts:
                found.extend(_normalize_weekday_tokens([part]))
            continue

        compact = re.sub(r"[^a-z]", "", text)
        if compact in mapping:
            found.extend(mapping[compact])
            continue

        if compact in ["mtwrf", "mtrf", "mwf"]:
            found.extend(_normalize_weekday_tokens([compact]))
            continue

        char_map = {
            "m": "Monday",
            "t": "Tuesday",
            "w": "Wednesday",
            "r": "Thursday",
            "f": "Friday",
            "u": "Sunday",
        }
        if len(compact) > 1 and compact.isalpha() and compact not in mapping:
            for char in compact:
                day = char_map.get(char)
                if day:
                    found.append(day)

    result = []
    for day in ordered_days:
        if day in found and day not in result:
            result.append(day)
    return result


# Returns the first plausible non-empty lines from a text block.
def _meaningful_lines(text, limit=20):
    lines = []
    for raw_line in (text or "").splitlines():
        line = _clean_space(raw_line)
        if not line:
            continue
        lines.append(line)
        if len(lines) >= limit:
            break
    return lines


# Returns a default year guess from the term start date.
def _default_year_from_term(course_defaults):
    term_start = course_defaults.get("term_start_date", "")
    if term_start:
        try:
            return int(term_start.split("-")[0])
        except Exception:
            pass
    return datetime.now().year


# Flattens parsed tables into row strings.
def _table_row_strings(tables) -> list:
    rows = []
    for table in tables or []:
        table_rows = table.get("rows", table)
        for row in table_rows or []:
            if isinstance(row, dict):
                row_text = " | ".join([_clean_space(str(value)) for value in row.values() if _clean_space(str(value))])
            else:
                row_text = " | ".join([_clean_space(str(cell)) for cell in row if _clean_space(str(cell))])
            if row_text:
                rows.append(row_text)
    return rows


# Builds a combined list of evidence lines from text blocks and tables.
def _evidence_lines(text, tables=None, parsed_document=None, limit=500):
    lines = []
    seen = set()

    for line in _meaningful_lines(text, limit=limit):
        key = line.lower()
        if key not in seen:
            seen.add(key)
            lines.append(line)

    for row_text in _table_row_strings(tables or []):
        key = row_text.lower()
        if key not in seen:
            seen.add(key)
            lines.append(row_text)

    if parsed_document:
        for block in parsed_document.get("blocks", []):
            block_text = _clean_space(block.get("text", ""))
            if not block_text:
                continue
            key = block_text.lower()
            if key not in seen:
                seen.add(key)
                lines.append(block_text)
            if len(lines) >= limit:
                break

    return lines[:limit]


# Deduplicates rows by a stable key.
def _dedupe_dict_rows(rows, key_fields):
    seen = set()
    output = []
    for row in rows:
        key = tuple((_clean_space(str(row.get(field, ""))) for field in key_fields))
        if key in seen:
            continue
        seen.add(key)
        output.append(row)
    return output


# Tries to pull the first matching date substring from a text line.
def _extract_first_date_in_text(text, default_year=None):
    if not text:
        return ""

    patterns = [
        r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2}(?:,\s*\d{4})?\b",
        r"\b\d{1,2}[/-]\d{1,2}(?:[/-]\d{2,4})?\b",
        r"\b\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s*(?:\d{4})?\b",
    ]

    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return _parse_date_string(match.group(0), default_year=default_year)

    return ""


# Tries to extract one or two time values from text.
def _extract_time_range(text):
    if not text:
        return "", ""

    cleaned_text = _clean_space(text).replace("–", "-").replace("—", "-")

    range_match = re.search(
        r"(\d{1,2}(?::\d{2})?\s*(?:AM|PM|am|pm)?)\s*(?:-|to|until|through)\s*(\d{1,2}(?::\d{2})?\s*(?:AM|PM|am|pm)?)",
        cleaned_text,
        flags=re.IGNORECASE,
    )
    if range_match:
        left_raw = _clean_space(range_match.group(1))
        right_raw = _clean_space(range_match.group(2))

        left_has_meridiem = bool(re.search(r"\b(?:AM|PM)\b", left_raw, flags=re.IGNORECASE))
        right_has_meridiem = bool(re.search(r"\b(?:AM|PM)\b", right_raw, flags=re.IGNORECASE))

        if not left_has_meridiem and right_has_meridiem:
            right_meridiem = re.search(r"\b(AM|PM)\b", right_raw, flags=re.IGNORECASE).group(1)
            left_raw = f"{left_raw} {right_meridiem}"

        if left_has_meridiem and not right_has_meridiem:
            left_meridiem = re.search(r"\b(AM|PM)\b", left_raw, flags=re.IGNORECASE).group(1)
            right_raw = f"{right_raw} {left_meridiem}"

        start_time = _parse_time_string(left_raw)
        end_time = _parse_time_string(right_raw)
        return start_time, end_time

    time_match = re.search(
        r"\b(\d{1,2}(?::\d{2})?\s*(?:AM|PM|am|pm)|[01]?\d:[0-5]\d|2[0-3]:[0-5]\d)\b",
        cleaned_text,
        flags=re.IGNORECASE,
    )
    if time_match:
        return _parse_time_string(time_match.group(1)), ""

    return "", ""

def _choose_session_title(primary_text, secondary_text):
    primary = _clean_space(primary_text)
    secondary = _clean_space(secondary_text)

    if re.fullmatch(r"(?:week|lecture|session)\s*\d+", primary, flags=re.IGNORECASE) and secondary:
        return secondary, primary

    return primary or secondary, secondary if primary and secondary and secondary != primary else ""


# Builds a confidence label from a numeric score.
def _confidence_label(score):
    if score >= 0.8:
        return "high"
    if score >= 0.5:
        return "medium"
    return "low"


# Infers recurring session dates from the term window and meeting days.
def _infer_session_dates(days_of_week, term_start, term_end, count) -> list:
    if not days_of_week or not term_start or not term_end or not count:
        return []

    day_numbers = {
        "Monday": 0,
        "Tuesday": 1,
        "Wednesday": 2,
        "Thursday": 3,
        "Friday": 4,
        "Saturday": 5,
        "Sunday": 6,
    }

    try:
        start_date = datetime.strptime(term_start, "%Y-%m-%d").date()
        end_date = datetime.strptime(term_end, "%Y-%m-%d").date()
    except ValueError:
        return []

    if end_date < start_date:
        return []

    target_days = {day_numbers.get(day) for day in days_of_week if day in day_numbers}
    target_days.discard(None)
    if not target_days:
        return []

    dates = []
    current = start_date
    while current <= end_date and len(dates) < count:
        if current.weekday() in target_days:
            dates.append(current.strftime("%Y-%m-%d"))
        current += timedelta(days=1)

    return dates


# -----------------------------
# Deterministic extraction layer
# -----------------------------

# Tries to extract the course name from text and tables.
def _extract_course_name(text, tables=None, parsed_document=None):
    patterns = [
        r"course\s+title\s*[:\-]\s*(.+)",
        r"course\s+name\s*[:\-]\s*(.+)",
        r"course\s+code\s*[:\-]\s*([A-Za-z]{2,}\s*\d{2,}[A-Za-z]?\s*[-: ]\s*.+)",
        r"\b([A-Za-z]{2,}\s*\d{2,}[A-Za-z]?)\b\s*[-:|]\s*([^\n]+)",
    ]

    for pattern in patterns:
        match = re.search(pattern, text or "", flags=re.IGNORECASE)
        if match:
            if len(match.groups()) == 1:
                return _clean_space(match.group(1))
            return _clean_space(" - ".join([part for part in match.groups() if part]))

    for row_text in _table_row_strings(tables or []):
        if re.search(r"course\s+(title|name|code)", row_text, flags=re.IGNORECASE):
            parts = [part.strip() for part in row_text.split("|") if part.strip()]
            if len(parts) >= 2:
                return _clean_space(" - ".join(parts[1:]))

    for line in _evidence_lines(text, tables=tables, parsed_document=parsed_document, limit=30):
        if re.search(r"course|syllabus|outline", line, flags=re.IGNORECASE):
            continue
        if len(line.split()) >= 2 and len(line) <= 120:
            return line

    return ""


# Extracts all normalized dates that appear in a text line.
def _extract_all_dates_in_text(text, default_year=None):
    if not text:
        return []

    matches = []
    patterns = [
        r"(?:[A-Za-z]{3,9}\.?\s+\d{1,2}(?:,\s*\d{4})?)",
        r"(?:\d{1,2}[/-]\d{1,2}(?:[/-]\d{2,4})?)",
        r"(?:\d{1,2}\s+[A-Za-z]{3,9}\.?\s*(?:\d{4})?)",
    ]

    for pattern in patterns:
        for match in re.finditer(pattern, text, flags=re.IGNORECASE):
            value = _parse_date_string(match.group(0), default_year=default_year)
            if value and value not in matches:
                matches.append(value)

    return matches


# Returns nearby lines after a label line such as DAY/TIME or COURSE FORMAT.
def _nearby_lines_after_label(lines, label_pattern, lookahead=4):
    for index, line in enumerate(lines):
        if re.search(label_pattern, line, flags=re.IGNORECASE):
            return lines[index + 1:index + 1 + lookahead]
    return []


# Returns a simple joined string for the non-empty tail of a table row.
def _join_table_tail_cells(row, start_index=2):
    if not isinstance(row, list):
        return ""
    return _clean_space(" ".join([cell for cell in row[start_index:] if _clean_space(cell)]))


# Detects a likely online delivery hint from text.
def _infer_location_from_delivery_context(text):
    text_lower = (text or "").lower()

    if re.search(r"\b(online|zoom|virtual|webinar|webinars|recorded)\b", text_lower):
        return "Online"

    room_match = re.search(
        r"\b(?:room|rm|classroom|venue|location)\s*[:\-]?\s*([A-Za-z0-9@._\- /]+)",
        text or "",
        flags=re.IGNORECASE,
    )
    if room_match:
        return _clean_space(room_match.group(1))

    return ""


# Extracts structured module plan rows from the course plan table when available.
def _extract_course_plan_rows(tables, course_defaults):
    plan_rows = []
    raw_lines = []

    for table in tables or []:
        table_rows = table.get("rows", table)
        if not table_rows or len(table_rows) < 3:
            continue

        header_preview = " ".join([" ".join(row) for row in table_rows[:2]])
        if not re.search(r"module", header_preview, flags=re.IGNORECASE):
            continue
        if not re.search(r"topic|lesson|webinar", header_preview, flags=re.IGNORECASE):
            continue

        for row in table_rows[2:]:
            if not isinstance(row, list):
                continue

            cells = [_clean_space(cell) for cell in row]
            if not any(cells):
                continue

            module_match = re.match(r"^(\d{1,2})$", cells[0] or "")
            if not module_match:
                continue

            module_number = int(module_match.group(1))
            topic = _clean_space(cells[1] if len(cells) > 1 else "")
            activity_text = _join_table_tail_cells(cells, start_index=2)
            raw_line = " | ".join([cell for cell in cells if cell])
            raw_lines.append(raw_line)

            plan_rows.append({
                "module_number": module_number,
                "topic": topic,
                "activity_text": activity_text,
                "raw_text": raw_line,
                "page_number": table.get("page_number", 0),
            })

    plan_rows = sorted(plan_rows, key=lambda item: item.get("module_number", 0))

    if plan_rows:
        inferred_dates = _infer_session_dates(
            course_defaults.get("days_of_week", []),
            course_defaults.get("term_start_date", ""),
            course_defaults.get("term_end_date", ""),
            len(plan_rows),
        )
        for index, row in enumerate(plan_rows):
            if index < len(inferred_dates):
                row["date"] = inferred_dates[index]
            else:
                row["date"] = ""
            row["date_inferred"] = bool(row.get("date"))
            row["start_time"] = course_defaults.get("start_time", "")
            row["end_time"] = course_defaults.get("end_time", "")

    return plan_rows, raw_lines


# Converts course plan rows into dated session records.
def _build_sessions_from_course_plan_rows(plan_rows, course_defaults):
    sessions = []
    raw_lines = []

    for row in plan_rows or []:
        topic = _clean_space(row.get("topic", ""))
        if not topic:
            continue

        activity_text = _clean_space(row.get("activity_text", ""))
        chunks = [part.strip("•∙ ").strip() for part in re.split(r"[•∙]+", activity_text) if part.strip()]
        session_chunks = []

        for chunk in chunks:
            chunk_lower = chunk.lower()
            if re.search(r"\b(assignment\s+\d+\s+(?:assigned|due)|term\s+project\s+due|project\s+due|midterm|final\s+exam|quiz\s+\d+)\b", chunk_lower, flags=re.IGNORECASE):
                continue
            session_chunks.append(chunk)

        session_description = " | ".join(session_chunks) if session_chunks else activity_text

        sessions.append({
            "session_number": int(row.get("module_number") or 0),
            "date": row.get("date", ""),
            "date_inferred": bool(row.get("date_inferred", False)),
            "start_time": row.get("start_time") or course_defaults.get("start_time", ""),
            "end_time": row.get("end_time") or course_defaults.get("end_time", ""),
            "title": topic,
            "description": session_description,
            "reminders_minutes": [60],
        })
        if row.get("raw_text"):
            raw_lines.append(row.get("raw_text"))

    return sessions, raw_lines

def _build_assignments_from_course_plan_rows(plan_rows, course_defaults):
    assignments = []
    raw_lines = []

    for row in plan_rows or []:
        activity_text = _clean_space(row.get("activity_text", ""))
        session_date = row.get("date", "")
        if not activity_text:
            continue

        chunks = [part.strip("•∙ ").strip() for part in re.split(r"[•∙]+", activity_text) if part.strip()]
        if not chunks:
            chunks = [activity_text]

        for chunk in chunks:
            chunk_clean = _clean_space(chunk)
            chunk_lower = chunk_clean.lower()

            if not chunk_clean:
                continue

            title = ""
            if re.search(r"assignment\s+\d+\s+due", chunk_clean, flags=re.IGNORECASE):
                match = re.search(r"(assignment\s+\d+)\s+due", chunk_clean, flags=re.IGNORECASE)
                title = match.group(1).title() if match else chunk_clean
            elif re.search(r"term\s+project\s+due", chunk_clean, flags=re.IGNORECASE):
                title = "Term Project"
            elif re.search(r"project\s+due", chunk_clean, flags=re.IGNORECASE):
                title = "Project"
            elif re.search(r"midterm", chunk_clean, flags=re.IGNORECASE):
                title = "Midterm"
            elif re.search(r"final\s+exam", chunk_clean, flags=re.IGNORECASE):
                title = "Final Exam"
            elif re.search(r"quiz\s+\d+", chunk_clean, flags=re.IGNORECASE) and re.search(r"due|test|quiz", chunk_clean, flags=re.IGNORECASE):
                match = re.search(r"(quiz\s+\d+)", chunk_clean, flags=re.IGNORECASE)
                title = match.group(1).title() if match else chunk_clean

            if not title:
                continue

            due_time, due_time_inferred = _extract_due_time(chunk_clean, course_defaults)
            if not due_time:
                due_time = "23:59"
                due_time_inferred = True

            assignments.append({
                "title": title,
                "due_date": session_date,
                "due_time": due_time,
                "due_time_inferred": due_time_inferred,
                "description": chunk_clean,
                "reminders_minutes": [10080, 1440],
            })
            raw_lines.append(f"Module {row.get('module_number')}: {chunk_clean}")

    return assignments, raw_lines

# Tries to extract the weekly schedule context from text and tables.
def _extract_schedule_details(text, tables=None, parsed_document=None, default_year=None):
    details = {
        "days_of_week": [],
        "start_time": "",
        "end_time": "",
        "term_start_date": "",
        "term_end_date": "",
        "location": "",
        "schedule_context": "",
        "term_context": "",
    }

    lines = _evidence_lines(text, tables=tables, parsed_document=parsed_document, limit=250)
    schedule_candidates = []
    term_candidates = []

    for line in lines:
        if re.search(r"\b(?:tuesday|thursday|monday|wednesday|friday|saturday|sunday|tue|thu|mon|wed|fri|mwf|tr|tth)\b", line, flags=re.IGNORECASE):
            schedule_candidates.append(line)

        all_dates = _extract_all_dates_in_text(line, default_year=default_year)
        if len(all_dates) >= 2:
            term_candidates.append((line, all_dates))

    for nearby_line in _nearby_lines_after_label(lines, r"^day/time:?$|\bday/time\b", lookahead=4):
        if nearby_line not in schedule_candidates and re.search(r"\b(?:tuesday|thursday|monday|wednesday|friday|saturday|sunday|tue|thu|mon|wed|fri|mwf|tr|tth)\b", nearby_line, flags=re.IGNORECASE):
            schedule_candidates.insert(0, nearby_line)

        all_dates = _extract_all_dates_in_text(nearby_line, default_year=default_year)
        if len(all_dates) >= 2:
            term_candidates.insert(0, (nearby_line, all_dates))

    for row_text in _table_row_strings(tables or []):
        if re.search(r"webinar", row_text, flags=re.IGNORECASE):
            schedule_candidates.append(row_text)
        all_dates = _extract_all_dates_in_text(row_text, default_year=default_year)
        if len(all_dates) >= 2:
            term_candidates.append((row_text, all_dates))

    chosen_schedule = ""
    for candidate in schedule_candidates:
        start_time, end_time = _extract_time_range(candidate)
        days = _normalize_weekday_tokens(
            re.findall(
                r"\b(?:MWF|TR|MW|TTh|Mon(?:day)?|Tue(?:s|sday)?|Wed(?:nesday)?|Thu(?:rs|rsday)?|Fri(?:day)?|Sat(?:urday)?|Sun(?:day)?|M|T|W|R|F)\b",
                candidate,
                flags=re.IGNORECASE,
            )
        )
        if days or start_time or end_time:
            chosen_schedule = candidate
            details["schedule_context"] = candidate
            details["days_of_week"] = days
            details["start_time"] = start_time
            details["end_time"] = end_time
            break

    for candidate_line, all_dates in term_candidates:
        if len(all_dates) >= 2:
            details["term_start_date"] = all_dates[0]
            details["term_end_date"] = all_dates[1]
            details["term_context"] = candidate_line
            break

    if not details["start_time"] or not details["end_time"]:
        for candidate in schedule_candidates:
            start_time, end_time = _extract_time_range(candidate)
            if start_time and end_time:
                details["start_time"] = details["start_time"] or start_time
                details["end_time"] = details["end_time"] or end_time
                if not details["schedule_context"]:
                    details["schedule_context"] = candidate
                break

    if not details["term_start_date"] or not details["term_end_date"]:
        begin_match = re.search(r"(?:classes\s+begin|start\s+date|term\s+start)\s*[:\-]\s*([^\n]+)", text or "", flags=re.IGNORECASE)
        end_match = re.search(r"(?:classes\s+end|end\s+date|term\s+end|last\s+day\s+of\s+class)\s*[:\-]\s*([^\n]+)", text or "", flags=re.IGNORECASE)
        if begin_match:
            details["term_start_date"] = _parse_date_string(begin_match.group(1), default_year=default_year)
        if end_match:
            default_end_year = default_year
            if details["term_start_date"]:
                default_end_year = int(details["term_start_date"].split("-")[0])
            details["term_end_date"] = _parse_date_string(end_match.group(1), default_year=default_end_year)

    location = ""
    if chosen_schedule:
        location = _infer_location_from_delivery_context(chosen_schedule)

    if not location:
        for row_text in _table_row_strings(tables or []):
            inferred_location = _infer_location_from_delivery_context(row_text)
            if inferred_location:
                location = inferred_location
                break

    if not location:
        location_match = re.search(
            r"(?:location|room|classroom|venue)\s*[:\-]\s*([^\n]+)",
            text or "",
            flags=re.IGNORECASE,
        )
        if location_match:
            location = _clean_space(location_match.group(1))

    details["location"] = location
    return details

def _extract_sessions_from_tables(tables, course_defaults):
    plan_rows, raw_lines = _extract_course_plan_rows(tables, course_defaults)
    if plan_rows:
        return _build_sessions_from_course_plan_rows(plan_rows, course_defaults)

    sessions = []
    raw_lines = []

    if not tables:
        return sessions, raw_lines

    for table in tables:
        table_rows = table.get("rows", table)
        if not table_rows:
            continue

        headers = [_clean_space(cell).lower() for cell in table_rows[0]]
        header_text = " | ".join(headers)
        if not re.search(r"week|session|lecture|date|topic|title|content", header_text, flags=re.IGNORECASE):
            continue

        column_map = {}
        for index, header in enumerate(headers):
            if re.search(r"week|session|lecture", header, flags=re.IGNORECASE):
                column_map.setdefault("session", index)
            if re.search(r"date", header, flags=re.IGNORECASE):
                column_map.setdefault("date", index)
            if re.search(r"topic|title|content|theme", header, flags=re.IGNORECASE):
                column_map.setdefault("topic", index)
            if re.search(r"description|details|notes", header, flags=re.IGNORECASE):
                column_map.setdefault("description", index)
            if re.search(r"time", header, flags=re.IGNORECASE):
                column_map.setdefault("time", index)

        for row in table_rows[1:]:
            if not any(_clean_space(cell) for cell in row):
                continue

            raw_lines.append(" | ".join(row))

            session_number = 0
            session_text = row[column_map["session"]] if "session" in column_map and column_map["session"] < len(row) else ""
            session_match = re.search(r"(\d+)", session_text)
            if session_match:
                session_number = int(session_match.group(1))

            date_text = row[column_map["date"]] if "date" in column_map and column_map["date"] < len(row) else ""
            date_value = _extract_first_date_in_text(date_text, default_year=_default_year_from_term(course_defaults))
            topic_text = row[column_map["topic"]] if "topic" in column_map and column_map["topic"] < len(row) else ""
            description_text = row[column_map["description"]] if "description" in column_map and column_map["description"] < len(row) else ""
            title, inherited_description = _choose_session_title(session_text, topic_text)
            if inherited_description and not description_text:
                description_text = inherited_description

            start_time = course_defaults.get("start_time", "")
            end_time = course_defaults.get("end_time", "")
            if "time" in column_map and column_map["time"] < len(row):
                row_start, row_end = _extract_time_range(row[column_map["time"]])
                start_time = row_start or start_time
                end_time = row_end or end_time

            if not title and description_text:
                title = description_text
                description_text = ""

            if title or date_value or session_number:
                sessions.append({
                    "session_number": session_number,
                    "date": date_value,
                    "date_inferred": False,
                    "start_time": start_time,
                    "end_time": end_time,
                    "title": _clean_space(title),
                    "description": _clean_space(description_text),
                    "reminders_minutes": [60],
                })

    return sessions, raw_lines

def _extract_sessions_from_text(text, course_defaults, parsed_document=None):
    sessions = []
    raw_lines = []
    default_year = _default_year_from_term(course_defaults)

    for line in _evidence_lines(text, parsed_document=parsed_document, limit=400):
        if not re.search(r"\b(week|session|lecture)\b", line, flags=re.IGNORECASE):
            continue

        if len(line) > 280:
            continue

        raw_lines.append(line)
        session_match = re.search(r"\b(?:week|session|lecture)\s*(\d+)\b", line, flags=re.IGNORECASE)
        session_number = int(session_match.group(1)) if session_match else 0

        date_value = _extract_first_date_in_text(line, default_year=default_year)
        start_time, end_time = _extract_time_range(line)
        title = _clean_space(line)
        description = ""

        if " - " in line:
            left, right = line.split(" - ", 1)
            title, inherited = _choose_session_title(left, right)
            description = inherited if inherited else right if right != title else ""

        if ":" in line and re.search(r"\b(week|session|lecture)\b", line, flags=re.IGNORECASE):
            left, right = line.split(":", 1)
            better_title, inherited = _choose_session_title(left, right)
            if better_title:
                title = better_title
                description = inherited if inherited else right

        sessions.append({
            "session_number": session_number,
            "date": date_value,
            "date_inferred": False,
            "start_time": start_time or course_defaults.get("start_time", ""),
            "end_time": end_time or course_defaults.get("end_time", ""),
            "title": _clean_space(title),
            "description": _clean_space(description),
            "reminders_minutes": [60],
        })

    return sessions, raw_lines

def _extract_due_time(text, course_defaults):
    if not text:
        return "", False

    if re.search(r"before\s+class", text, flags=re.IGNORECASE):
        class_start = course_defaults.get("start_time", "")
        return class_start, bool(class_start)

    if re.search(r"\b(midnight|end of day|eod)\b", text, flags=re.IGNORECASE):
        return "23:59", True

    time_value = _extract_time_range(text)[0]
    if time_value:
        return time_value, False

    return "", False


# Extracts assignment and assessment rows from native tables.
def _extract_assignments_from_tables(tables, course_defaults):
    plan_rows, raw_lines = _extract_course_plan_rows(tables, course_defaults)
    if plan_rows:
        return _build_assignments_from_course_plan_rows(plan_rows, course_defaults)

    assignments = []
    raw_lines = []
    default_year = _default_year_from_term(course_defaults)

    if not tables:
        return assignments, raw_lines

    for table in tables:
        table_rows = table.get("rows", table)
        if not table_rows:
            continue

        headers = [_clean_space(cell).lower() for cell in table_rows[0]]
        header_text = " | ".join(headers)
        if not re.search(r"assignment|quiz|exam|project|deadline|due", header_text, flags=re.IGNORECASE):
            continue

        column_map = {}
        for index, header in enumerate(headers):
            if re.search(r"assignment|quiz|exam|project|assessment|task|title", header, flags=re.IGNORECASE):
                column_map.setdefault("title", index)
            if re.search(r"due|deadline|date", header, flags=re.IGNORECASE):
                column_map.setdefault("date", index)
            if re.search(r"time", header, flags=re.IGNORECASE):
                column_map.setdefault("time", index)
            if re.search(r"description|details|notes", header, flags=re.IGNORECASE):
                column_map.setdefault("description", index)

        for row in table_rows[1:]:
            if not any(_clean_space(cell) for cell in row):
                continue

            row_text = " | ".join(row)
            if not re.search(r"\b(due|midterm|final\s+exam|quiz\s+\d+|project\s+due|assignment\s+\d+\s+due)\b", row_text, flags=re.IGNORECASE):
                continue

            raw_lines.append(row_text)
            title = row[column_map["title"]] if "title" in column_map and column_map["title"] < len(row) else row_text
            due_source = row[column_map["date"]] if "date" in column_map and column_map["date"] < len(row) else row_text
            due_date = _extract_first_date_in_text(due_source, default_year=default_year)

            time_source = row[column_map["time"]] if "time" in column_map and column_map["time"] < len(row) else row_text
            due_time, due_inferred = _extract_due_time(time_source, course_defaults)
            description = row[column_map["description"]] if "description" in column_map and column_map["description"] < len(row) else ""

            if not due_time:
                due_time = "23:59"
                due_inferred = True

            assignments.append({
                "title": _clean_space(title),
                "due_date": due_date,
                "due_time": due_time,
                "due_time_inferred": due_inferred,
                "description": _clean_space(description),
                "reminders_minutes": [10080, 1440],
            })

    return assignments, raw_lines

def _extract_assignments_from_text(text, course_defaults, parsed_document=None):
    assignments = []
    raw_lines = []
    default_year = _default_year_from_term(course_defaults)

    for line in _evidence_lines(text, parsed_document=parsed_document, limit=500):
        if not re.search(r"\b(due|midterm|final\s+exam|quiz\s+\d+|project\s+due|assignment\s+\d+\s+due)\b", line, flags=re.IGNORECASE):
            continue

        if len(line) > 280:
            continue

        raw_lines.append(line)
        due_date = _extract_first_date_in_text(line, default_year=default_year)
        due_time, due_inferred = _extract_due_time(line, course_defaults)
        if not due_time:
            due_time = "23:59"
            due_inferred = True

        title = line
        description = ""

        if ":" in line:
            left, right = line.split(":", 1)
            if re.search(r"\b(assignment|quiz|exam|project|midterm|final)\b", left, flags=re.IGNORECASE):
                title = left
                description = right
            else:
                description = line

        assignments.append({
            "title": _clean_space(title),
            "due_date": due_date,
            "due_time": due_time,
            "due_time_inferred": due_inferred,
            "description": _clean_space(description),
            "reminders_minutes": [10080, 1440],
        })

    return assignments, raw_lines

def _fill_missing_session_dates(sessions, course_defaults):
    if not sessions:
        return sessions

    missing = [row for row in sessions if not row.get("date")]
    if not missing:
        return sessions

    inferred_dates = _infer_session_dates(
        course_defaults.get("days_of_week", []),
        course_defaults.get("term_start_date", ""),
        course_defaults.get("term_end_date", ""),
        len(sessions),
    )

    if not inferred_dates:
        return sessions

    sorted_sessions = sorted(
        sessions,
        key=lambda item: (
            item.get("session_number", 999999) if item.get("session_number") else 999999,
            item.get("title", ""),
        ),
    )

    inferred_index = 0
    for row in sorted_sessions:
        if not row.get("date") and inferred_index < len(inferred_dates):
            row["date"] = inferred_dates[inferred_index]
            row["date_inferred"] = True
            inferred_index += 1

    return sorted_sessions


# Normalizes and fills gaps in session rows.
def _finalize_sessions(sessions, course_defaults):
    output = []
    for row in sessions:
        if not isinstance(row, dict):
            continue
        title = _clean_space(row.get("title", ""))
        description = _clean_space(row.get("description", ""))
        if not title and not description and not row.get("date"):
            continue

        start_time = row.get("start_time") or course_defaults.get("start_time", "")
        end_time = row.get("end_time") or course_defaults.get("end_time", "")
        if start_time and not end_time:
            end_time = _infer_end_time(start_time)

        output.append({
            "session_number": int(row.get("session_number") or 0),
            "date": row.get("date", ""),
            "date_inferred": bool(row.get("date_inferred", False)),
            "start_time": start_time,
            "end_time": end_time,
            "title": title or description or f"Class session {len(output) + 1}",
            "description": description,
            "reminders_minutes": [60],
        })

    output = _fill_missing_session_dates(output, course_defaults)
    output = _dedupe_dict_rows(output, ["session_number", "date", "title"])
    return output


# Normalizes and fills gaps in assignment rows.
def _finalize_assignments(assignments):
    output = []
    for row in assignments:
        if not isinstance(row, dict):
            continue
        title = _clean_space(row.get("title", ""))
        due_date = row.get("due_date", "")
        description = _clean_space(row.get("description", ""))
        due_time = row.get("due_time", "") or ""
        due_time_inferred = bool(row.get("due_time_inferred", False))

        if due_date and not due_time:
            due_time = "23:59"
            due_time_inferred = True

        if not title and description:
            title = description

        if not title and not due_date:
            continue

        output.append({
            "title": title or description or "Course assessment",
            "due_date": due_date,
            "due_time": due_time,
            "due_time_inferred": due_time_inferred,
            "description": description,
            "reminders_minutes": [10080, 1440],
        })

    output = _dedupe_dict_rows(output, ["title", "due_date", "due_time"])
    return output



# -----------------------------
# Validation layer
# -----------------------------

# Validates and normalizes course outline data into the UI-compatible schema.
def validate_course_outline_data(data: dict) -> dict:
    default_result = {
        "course_name": "",
        "days_of_week": [],
        "start_time": "",
        "end_time": "",
        "term_start_date": "",
        "term_end_date": "",
        "location": "",
        "sessions": [],
        "assignments": [],
        "confidence": 0.0,
        "confidence_label": "low",
        "extraction_contexts": {
            "course_name_context": "",
            "schedule_context": "",
            "sessions_raw": "",
            "assignments_raw": "",
            "term_context": "",
        },
        "parse_meta": {},
        "extraction_meta": {},
    }

    source = data if isinstance(data, dict) else {}
    result = dict(default_result)

    result["course_name"] = _clean_space(source.get("course_name", ""))
    result["days_of_week"] = _normalize_weekday_tokens(source.get("days_of_week", []))
    result["start_time"] = _parse_time_string(source.get("start_time", "")) or ""
    result["end_time"] = _parse_time_string(source.get("end_time", "")) or ""
    result["term_start_date"] = _parse_date_string(source.get("term_start_date", ""))
    term_end_default_year = int(result["term_start_date"].split("-")[0]) if result["term_start_date"] else None
    result["term_end_date"] = _parse_date_string(source.get("term_end_date", ""), default_year=term_end_default_year)
    result["location"] = _clean_space(source.get("location", ""))

    if result["start_time"] and not result["end_time"]:
        result["end_time"] = _infer_end_time(result["start_time"])

    if result["term_start_date"] and result["term_end_date"] and result["term_end_date"] < result["term_start_date"]:
        result["term_start_date"], result["term_end_date"] = result["term_end_date"], result["term_start_date"]

    course_defaults = {
        "days_of_week": result["days_of_week"],
        "start_time": result["start_time"],
        "end_time": result["end_time"],
        "term_start_date": result["term_start_date"],
        "term_end_date": result["term_end_date"],
        "location": result["location"],
    }

    sessions = []
    for row in source.get("sessions", []) or []:
        if not isinstance(row, dict):
            continue

        title = _clean_space(row.get("title", ""))
        description = _clean_space(row.get("description", ""))
        date_value = _parse_date_string(row.get("date", ""), default_year=_default_year_from_term(course_defaults))
        start_time = _parse_time_string(row.get("start_time", "")) or course_defaults.get("start_time", "")
        end_time = _parse_time_string(row.get("end_time", "")) or course_defaults.get("end_time", "")

        if start_time and not end_time:
            end_time = _infer_end_time(start_time)

        try:
            session_number = int(row.get("session_number") or 0)
        except Exception:
            session_number = 0

        if not title and description:
            title = description
            description = ""

        if not title and not date_value:
            continue

        sessions.append({
            "session_number": session_number,
            "date": date_value,
            "date_inferred": bool(row.get("date_inferred", False)),
            "start_time": start_time,
            "end_time": end_time,
            "title": title or f"Class session {len(sessions) + 1}",
            "description": description,
            "reminders_minutes": [60],
        })

    assignments = []
    for row in source.get("assignments", []) or []:
        if not isinstance(row, dict):
            continue

        title = _clean_space(row.get("title", ""))
        description = _clean_space(row.get("description", ""))
        due_date = _parse_date_string(row.get("due_date", ""), default_year=_default_year_from_term(course_defaults))
        due_time = _parse_time_string(row.get("due_time", "")) or ""

        due_time_inferred = bool(row.get("due_time_inferred", False))
        if due_date and not due_time:
            due_time = "23:59"
            due_time_inferred = True

        if not title and description:
            title = description
            description = ""

        if not title and not due_date:
            continue

        assignments.append({
            "title": title or "Course assessment",
            "due_date": due_date,
            "due_time": due_time,
            "due_time_inferred": due_time_inferred,
            "description": description,
            "reminders_minutes": [10080, 1440],
        })

    result["sessions"] = _fill_missing_session_dates(_dedupe_dict_rows(sessions, ["session_number", "date", "title"]), course_defaults)
    result["assignments"] = _dedupe_dict_rows(assignments, ["title", "due_date", "due_time"])

    confidence = source.get("confidence", 0.0)
    try:
        confidence_value = float(confidence)
    except Exception:
        confidence_value = 0.0
    confidence_value = max(0.0, min(confidence_value, 1.0))
    result["confidence"] = round(confidence_value, 2)
    result["confidence_label"] = _confidence_label(result["confidence"])

    extraction_contexts = source.get("extraction_contexts", {}) or {}
    if isinstance(extraction_contexts, dict):
        merged_contexts = dict(default_result["extraction_contexts"])
        for key, value in extraction_contexts.items():
            merged_contexts[key] = value
        result["extraction_contexts"] = merged_contexts

    parse_meta = source.get("parse_meta", {}) or {}
    if isinstance(parse_meta, dict):
        result["parse_meta"] = parse_meta

    extraction_meta = source.get("extraction_meta", {}) or {}
    if isinstance(extraction_meta, dict):
        result["extraction_meta"] = extraction_meta

    if not result["confidence"]:
        score = 0.0
        if result["course_name"]:
            score += 0.15
        if result["days_of_week"]:
            score += 0.15
        if result["start_time"] and result["end_time"]:
            score += 0.1
        if result["term_start_date"] and result["term_end_date"]:
            score += 0.15
        if result["location"]:
            score += 0.05
        if result["sessions"]:
            score += 0.25
        if result["assignments"]:
            score += 0.15
        result["confidence"] = round(min(score, 1.0), 2)
        result["confidence_label"] = _confidence_label(result["confidence"])

    return result


# Validates and normalizes invitation details into the UI-compatible schema.
def validate_invitation_details(data: dict) -> dict:
    default_result = {
        "event_title": "",
        "event_date": "",
        "event_time": "",
        "event_end_time": "",
        "location": "",
        "host": "",
        "rsvp_required": False,
        "rsvp_deadline": "",
        "description": "",
        "confidence": 0.0,
        "confidence_label": "low",
        "parse_meta": {},
        "extraction_meta": {},
    }

    source = data if isinstance(data, dict) else {}
    result = dict(default_result)

    result["event_title"] = _clean_space(source.get("event_title", ""))
    result["event_date"] = _parse_date_string(source.get("event_date", ""))
    result["event_time"] = _parse_time_string(source.get("event_time", "")) or ""
    result["event_end_time"] = _parse_time_string(source.get("event_end_time", "")) or ""
    result["location"] = _clean_space(source.get("location", ""))
    result["host"] = _clean_space(source.get("host", ""))
    result["rsvp_required"] = bool(source.get("rsvp_required", False))
    default_rsvp_year = int(result["event_date"].split("-")[0]) if result["event_date"] else None
    result["rsvp_deadline"] = _parse_date_string(source.get("rsvp_deadline", ""), default_year=default_rsvp_year)
    result["description"] = _clean_space(source.get("description", ""))

    if result["event_time"] and not result["event_end_time"]:
        result["event_end_time"] = _infer_end_time(result["event_time"])

    confidence = source.get("confidence", 0.0)
    try:
        confidence_value = float(confidence)
    except Exception:
        confidence_value = 0.0
    confidence_value = max(0.0, min(confidence_value, 1.0))
    result["confidence"] = round(confidence_value, 2)
    result["confidence_label"] = _confidence_label(result["confidence"])

    if not result["confidence"]:
        score = 0.0
        if result["event_title"]:
            score += 0.25
        if result["event_date"]:
            score += 0.25
        if result["event_time"]:
            score += 0.15
        if result["location"]:
            score += 0.15
        if result["description"]:
            score += 0.1
        if result["rsvp_required"]:
            score += 0.1
        result["confidence"] = round(min(score, 1.0), 2)
        result["confidence_label"] = _confidence_label(result["confidence"])

    parse_meta = source.get("parse_meta", {}) or {}
    if isinstance(parse_meta, dict):
        result["parse_meta"] = parse_meta

    extraction_meta = source.get("extraction_meta", {}) or {}
    if isinstance(extraction_meta, dict):
        result["extraction_meta"] = extraction_meta

    return result


# Extracts a structured course outline object using deterministic parsing and normalization.
def extract_course_outline(text: str, tables=None, parsed_document=None) -> dict:
    base = {
        "course_name": "",
        "days_of_week": [],
        "start_time": "",
        "end_time": "",
        "term_start_date": "",
        "term_end_date": "",
        "location": "",
        "sessions": [],
        "assignments": [],
        "confidence": 0.0,
        "confidence_label": "low",
        "extraction_contexts": {
            "course_name_context": "",
            "schedule_context": "",
            "sessions_raw": "",
            "assignments_raw": "",
            "term_context": "",
            "course_plan_raw": "",
        },
        "parse_meta": {},
        "extraction_meta": {
            "provider": "deterministic",
            "status": "success",
            "message": "Deterministic parser result.",
        },
    }

    if not text and parsed_document:
        text = parsed_document.get("raw_text", "")

    if not text:
        return base

    try:
        tables = tables if tables is not None else (parsed_document.get("tables", []) if parsed_document else [])
        course_name = _extract_course_name(text, tables=tables, parsed_document=parsed_document)
        schedule_details = _extract_schedule_details(
            text,
            tables=tables,
            parsed_document=parsed_document,
            default_year=datetime.now().year,
        )

        course_defaults = {
            "days_of_week": schedule_details.get("days_of_week", []),
            "start_time": schedule_details.get("start_time", ""),
            "end_time": schedule_details.get("end_time", ""),
            "term_start_date": schedule_details.get("term_start_date", ""),
            "term_end_date": schedule_details.get("term_end_date", ""),
            "location": schedule_details.get("location", ""),
        }

        course_plan_rows, course_plan_raw_lines = _extract_course_plan_rows(tables or [], course_defaults)

        if course_plan_rows:
            sessions, sessions_raw_lines = _build_sessions_from_course_plan_rows(course_plan_rows, course_defaults)
            assignments, assignments_raw_lines = _build_assignments_from_course_plan_rows(course_plan_rows, course_defaults)
        else:
            session_rows_tables, sessions_raw_tables = _extract_sessions_from_tables(tables or [], course_defaults)
            session_rows_text, sessions_raw_text = _extract_sessions_from_text(text, course_defaults, parsed_document=parsed_document)
            sessions = _finalize_sessions(session_rows_tables + session_rows_text, course_defaults)

            assignment_rows_tables, assignments_raw_tables = _extract_assignments_from_tables(tables or [], course_defaults)
            assignment_rows_text, assignments_raw_text = _extract_assignments_from_text(text, course_defaults, parsed_document=parsed_document)
            assignments = _finalize_assignments(assignment_rows_tables + assignment_rows_text)
            sessions_raw_lines = sessions_raw_tables + sessions_raw_text
            assignments_raw_lines = assignments_raw_tables + assignments_raw_text

        result = {
            "course_name": course_name,
            "days_of_week": course_defaults["days_of_week"],
            "start_time": course_defaults["start_time"],
            "end_time": course_defaults["end_time"],
            "term_start_date": course_defaults["term_start_date"],
            "term_end_date": course_defaults["term_end_date"],
            "location": course_defaults["location"],
            "sessions": _finalize_sessions(sessions, course_defaults),
            "assignments": _finalize_assignments(assignments),
            "extraction_contexts": {
                "course_name_context": _get_context_window(text, course_name.split(" - ")[0] if course_name else "course") if course_name else _get_context_window(text, "course"),
                "schedule_context": schedule_details.get("schedule_context", ""),
                "sessions_raw": "\n".join((sessions_raw_lines or [])[:30]),
                "assignments_raw": "\n".join((assignments_raw_lines or [])[:30]),
                "term_context": schedule_details.get("term_context", ""),
                "course_plan_raw": "\n".join((course_plan_raw_lines or [])[:30]),
            },
            "parse_meta": parsed_document.get("quality", {}) if parsed_document else {},
            "extraction_meta": {
                "provider": "deterministic",
                "status": "success",
                "message": "Deterministic parser result.",
            },
        }

        return validate_course_outline_data(result)
    except Exception as e:
        base["extraction_meta"]["status"] = "error"
        base["extraction_meta"]["message"] = str(e)
        return base

def _extract_invitation_title(text):
    lines = _meaningful_lines(text, limit=20)
    generic_patterns = [
        r"you are invited",
        r"cordially invited",
        r"please join us",
        r"join us for",
        r"kindly rsvp",
        r"rsvp",
    ]

    phrase_match = re.search(r"(?:join us for|invite you to|celebration of)\s+(.+)", text or "", flags=re.IGNORECASE)
    if phrase_match:
        title = _clean_space(phrase_match.group(1).split("\n")[0])
        if title:
            return title

    for line in lines:
        if any(re.search(pattern, line, flags=re.IGNORECASE) for pattern in generic_patterns):
            continue
        if re.search(r"\b(date|time|venue|location|rsvp|hosted by|from)\b", line, flags=re.IGNORECASE):
            continue
        if 2 <= len(line.split()) <= 12:
            return line

    return lines[0] if lines else ""


# Extracts a short description from the first few meaningful lines.
def _build_short_description(text):
    meaningful = []
    for line in _meaningful_lines(text, limit=12):
        if re.search(r"\b(date|time|venue|location|rsvp)\b", line, flags=re.IGNORECASE):
            continue
        meaningful.append(line)
        if len(meaningful) >= 3:
            break
    return "\n".join(meaningful)


# Extracts a structured invitation object using deterministic parsing and normalization.
def extract_invitation_details(text: str, parsed_document=None, tables=None) -> dict:
    base = {
        "event_title": "",
        "event_date": "",
        "event_time": "",
        "event_end_time": "",
        "location": "",
        "host": "",
        "rsvp_required": False,
        "rsvp_deadline": "",
        "description": "",
        "confidence": 0.0,
        "confidence_label": "low",
        "parse_meta": {},
        "extraction_meta": {
            "provider": "deterministic",
            "status": "success",
            "message": "Deterministic parser result.",
        },
    }

    if not text and parsed_document:
        text = parsed_document.get("raw_text", "")

    if not text:
        return base

    try:
        lines = _evidence_lines(text, tables=tables, parsed_document=parsed_document, limit=120)
        event_title = _extract_invitation_title(text)
        default_year = datetime.now().year
        event_date = _extract_first_date_in_text(text, default_year=default_year)
        event_time, event_end_time = _extract_time_range(text)

        location = ""
        location_match = re.search(r"(?:location|venue|where|address)\s*[:\-]\s*([^\n]+)", text, flags=re.IGNORECASE)
        if location_match:
            location = _clean_space(location_match.group(1))
        else:
            for line in lines:
                if re.search(r"\b(venue|location|where|address)\b", line, flags=re.IGNORECASE):
                    cleaned_line = re.sub(r"^(?:venue|location|where|address)\s*[:\-]?\s*", "", line, flags=re.IGNORECASE)
                    if cleaned_line and cleaned_line != line:
                        location = _clean_space(cleaned_line)
                        break

        host = ""
        host_match = re.search(r"(?:hosted by|from|organizer|organised by|organized by)\s*[:\-]?\s*([^\n]+)", text, flags=re.IGNORECASE)
        if host_match:
            host = _clean_space(host_match.group(1))

        rsvp_required = bool(re.search(r"\brsvp\b", text, flags=re.IGNORECASE))
        rsvp_deadline = ""
        rsvp_match = re.search(r"(?:rsvp(?:\s+by)?|kindly rsvp by)\s*[:\-]?\s*([^\n]+)", text, flags=re.IGNORECASE)
        if rsvp_match:
            rsvp_deadline = _extract_first_date_in_text(rsvp_match.group(1), default_year=default_year)

        description = _build_short_description(text)

        result = {
            "event_title": event_title,
            "event_date": event_date,
            "event_time": event_time,
            "event_end_time": event_end_time,
            "location": location,
            "host": host,
            "rsvp_required": rsvp_required,
            "rsvp_deadline": rsvp_deadline,
            "description": description,
            "parse_meta": parsed_document.get("quality", {}) if parsed_document else {},
            "extraction_meta": {
                "provider": "deterministic",
                "status": "success",
                "message": "Deterministic parser result.",
            },
        }

        return validate_invitation_details(result)
    except Exception as e:
        base["extraction_meta"]["status"] = "error"
        base["extraction_meta"]["message"] = str(e)
        return base
